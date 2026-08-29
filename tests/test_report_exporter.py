"""合规审查报告 Word 导出单测。"""

from engine.review.report_exporter import build_compliance_report_docx, default_export_filename


def test_build_report_with_risks_and_human_note():
    report = {
        "material_id": "abc-123-def",
        "file_name": "demo.docx",
        "scene": "营销宣传",
        "summary": "删除保证收益与赠礼表述。",
        "risk_sentences": [
            {
                "text": "锁定12%稳赚收益",
                "risk_types": ["承诺收益", "保证收益"],
                "hit_case_id": "C001",
                "hit_party_name": "某某寿险",
                "hit_penalty_doc_no": "京保监罚〔2014〕29号",
                "case_key_field": "承诺收益类违法事实摘要",
                "match_reason": "语义与标签命中保证收益",
            }
        ],
    }
    data = build_compliance_report_docx(
        report, human_note="同意系统结论", human_review_done=True,
    )
    assert data[:2] == b"PK"
    assert "合规审查报告" in default_export_filename(report)


def test_human_note_absent_when_not_done():
    from io import BytesIO
    from docx import Document

    report = {
        "material_id": "x",
        "summary": "改写材料",
        "risk_sentences": [],
    }
    raw = build_compliance_report_docx(
        report, human_note="不应出现", human_review_done=False,
    )
    doc = Document(BytesIO(raw))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            parts.append(ncell.text)
    texts = "\n".join(parts)
    assert "三、人工复核建议" in texts
    assert "供工作人员填写复核意见" in texts
    assert "不应出现" not in texts

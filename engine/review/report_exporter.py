"""按赛题《合规审查报告模板》结构导出 Word；视觉对齐 Trust & Authority（法务深蓝）。"""

from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Legal / E-signature tokens（ui-ux-pro-max：Trust navy + 可读正文）
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x64, 0x74, 0x8B)
BORDER = "CBD5E1"
CARD_BG = "F8FAFC"
HEADER_BG = "1E3A5F"
LABEL_BG = "E8EEF5"
ACCENT_LINE = "2563EB"


def _set_run_font(
    run,
    *,
    size: int = 11,
    bold: bool = False,
    color: RGBColor | None = None,
    name: str = "微软雅黑",
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color is not None:
        run.font.color.rgb = color
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:eastAsia"), name)
    except Exception:  # noqa: BLE001
        pass


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = BORDER, sz: str = "8") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, *, top=60, bottom=60, left=100, right=100) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _clear_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _para_spacing(p, *, before=0, after=6, line: float | None = 1.15) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def _add_section_heading(doc: Document, text: str) -> None:
    """章节标题：左侧强调色竖条 + 深蓝字重。"""
    table = doc.add_table(rows=1, cols=2)
    _clear_table_borders(table)
    table.autofit = True
    bar, title = table.rows[0].cells
    bar.width = Cm(0.25)
    title.width = Cm(15.5)
    _shade_cell(bar, ACCENT_LINE)
    p = title.paragraphs[0]
    run = p.add_run(text)
    _set_run_font(run, size=14, bold=True, color=NAVY)
    _para_spacing(p, before=2, after=2)
    _set_cell_margins(title, top=40, bottom=40, left=120, right=40)
    _set_cell_margins(bar, top=0, bottom=0, left=0, right=0)
    spacer = doc.add_paragraph()
    _para_spacing(spacer, before=0, after=8)


def _add_body(doc: Document, text: str, *, muted: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text or "")
    _set_run_font(run, size=11, color=MUTED if muted else INK)
    _para_spacing(p, before=0, after=8)


def _add_header_banner(doc: Document, title: str, subtitle: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    _clear_table_borders(table)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, HEADER_BG)
    _set_cell_margins(cell, top=140, bottom=120, left=160, right=160)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_run_font(run, size=22, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), name="微软雅黑")
    _para_spacing(p, before=0, after=4)

    line = cell.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = line.add_run("◆  —————————  ◆")
    _set_run_font(lr, size=9, color=RGBColor(0x93, 0xC5, 0xFD))
    _para_spacing(line, before=2, after=2)

    if subtitle:
        sub = cell.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(subtitle)
        _set_run_font(sr, size=9, color=RGBColor(0xBF, 0xDB, 0xFE))
        _para_spacing(sub, before=2, after=0)

    gap = doc.add_paragraph()
    _para_spacing(gap, before=0, after=10)


def _add_meta_bar(doc: Document, left: str, right: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    _clear_table_borders(table)
    c0, c1 = table.rows[0].cells
    _shade_cell(c0, LABEL_BG)
    _shade_cell(c1, LABEL_BG)
    for cell, text, align in (
        (c0, left, WD_ALIGN_PARAGRAPH.LEFT),
        (c1, right, WD_ALIGN_PARAGRAPH.RIGHT),
    ):
        _set_cell_margins(cell, top=50, bottom=50, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = align
        if "：" in text:
            lab, val = text.split("：", 1)
            r1 = p.add_run(f"{lab}：")
            _set_run_font(r1, size=9, bold=True, color=MUTED)
            r2 = p.add_run(val)
            _set_run_font(r2, size=9, color=INK)
        else:
            r = p.add_run(text)
            _set_run_font(r, size=9, color=MUTED)
    gap = doc.add_paragraph()
    _para_spacing(gap, before=0, after=12)


def _add_kv_card(doc: Document, badge: str, rows: list[tuple[str, str]]) -> None:
    """风险点卡片：顶栏徽章 + 标签/内容表。"""
    wrap = doc.add_table(rows=1, cols=1)
    _clear_table_borders(wrap)
    outer = wrap.rows[0].cells[0]
    _shade_cell(outer, "FFFFFF")
    _set_cell_borders(outer, color="93C5FD", sz="12")
    _set_cell_margins(outer, top=80, bottom=80, left=100, right=100)

    badge_tbl = outer.add_table(rows=1, cols=1)
    _clear_table_borders(badge_tbl)
    bc = badge_tbl.rows[0].cells[0]
    _shade_cell(bc, HEADER_BG)
    _set_cell_margins(bc, top=30, bottom=30, left=80, right=80)
    bp2 = bc.paragraphs[0]
    br2 = bp2.add_run(badge)
    _set_run_font(br2, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # remove empty default paragraph gap under badge
    if outer.paragraphs and not outer.paragraphs[0].text:
        p0 = outer.paragraphs[0]
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(4)

    kv = outer.add_table(rows=len(rows), cols=2)
    _clear_table_borders(kv)
    for i, (label, value) in enumerate(rows):
        lc, vc = kv.rows[i].cells
        lc.width = Cm(2.8)
        vc.width = Cm(12.2)
        _set_cell_margins(lc, top=40, bottom=40, left=40, right=40)
        _set_cell_margins(vc, top=40, bottom=40, left=40, right=40)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        _set_run_font(lr, size=9, bold=True, color=MUTED)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value or "—")
        _set_run_font(vr, size=10, color=INK)
        _para_spacing(lp, before=0, after=0, line=1.15)
        _para_spacing(vp, before=0, after=0, line=1.15)
        if i < len(rows) - 1:
            for cell in (lc, vc):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "dashed")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:color"), "BFDBFE")
                borders.append(bottom)
                tcPr.append(borders)

    gap = doc.add_paragraph()
    _para_spacing(gap, before=4, after=10)


def _add_review_box(doc: Document, text: str, *, placeholder: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    _clear_table_borders(table)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, CARD_BG)
    _set_cell_borders(cell, color="93C5FD", sz="12")
    _set_cell_margins(cell, top=120, bottom=160, left=120, right=120)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _set_run_font(run, size=10, color=MUTED if placeholder else INK)
    _para_spacing(p, before=0, after=0, line=1.25)
    # min height feel
    for _ in range(2):
        blank = cell.add_paragraph()
        _para_spacing(blank, before=0, after=8)


def _material_source(report: dict[str, Any]) -> str:
    source_file = (
        report.get("file_name") or report.get("source_file") or ""
    ).strip()
    if source_file:
        return f"文件《{source_file}》"
    return "在线文本输入"


def _risk_tags(item: dict[str, Any]) -> str:
    tags = item.get("risk_types") or []
    if isinstance(tags, list) and tags:
        return "、".join(str(t) for t in tags if t)
    ids = item.get("risk_type_ids") or []
    if isinstance(ids, list) and ids:
        return "、".join(str(t) for t in ids if t)
    return "—"


def _hit_case_line(item: dict[str, Any]) -> str:
    cid = (item.get("hit_case_id") or "").strip()
    party = (item.get("hit_party_name") or "").strip()
    if cid and party:
        return f"{cid} · {party}"
    return cid or party or "—"


def _hit_doc_no(item: dict[str, Any]) -> str:
    return (item.get("hit_penalty_doc_no") or "").strip() or "—"


def _case_fact(item: dict[str, Any]) -> str:
    vb = (item.get("case_key_field") or "").strip()
    if not vb:
        cases = item.get("retrieved_cases") or []
        if cases and isinstance(cases[0], dict):
            vb = (cases[0].get("violation_behavior") or "")[:160]
    return vb or "—"


def _safe_filename(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*\s]+', "_", name).strip("._") or "report"
    return name[:80]


def build_compliance_report_docx(
    report: dict[str, Any],
    *,
    human_note: str | None = None,
    human_review_done: bool = False,
) -> bytes:
    """生成版式化合规审查报告（docx）。"""
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    _add_header_banner(doc, "合规审查报告", "保险营销材料 · 智能合规审查")
    _add_meta_bar(
        doc,
        f"材料来源：{_material_source(report)}",
        f"审查时间：{now}",
    )

    scene = (report.get("scene") or "").strip()
    if scene:
        _add_body(doc, f"审查场景：{scene}", muted=True)

    # 一、整改建议
    _add_section_heading(doc, "一、整改建议")
    suggestion = (
        (report.get("summary") or "").strip()
        or (report.get("overall_suggestion") or "").strip()
        or "无"
    )
    tip_box = doc.add_table(rows=1, cols=1)
    _clear_table_borders(tip_box)
    tip_cell = tip_box.rows[0].cells[0]
    _shade_cell(tip_cell, "EFF6FF")
    _set_cell_borders(tip_cell, color="BFDBFE", sz="8")
    _set_cell_margins(tip_cell, top=80, bottom=80, left=120, right=120)
    tp = tip_cell.paragraphs[0]
    tr = tp.add_run(suggestion)
    _set_run_font(tr, size=11, color=INK)
    _para_spacing(tp, before=0, after=0, line=1.25)
    gap = doc.add_paragraph()
    _para_spacing(gap, before=4, after=8)

    # 二、风险识别详情
    _add_section_heading(doc, "二、风险识别详情")
    risk_sentences = report.get("risk_sentences") or []
    if not risk_sentences:
        _add_body(doc, "未识别到风险语句。", muted=True)
    else:
        for i, item in enumerate(risk_sentences, 1):
            if not isinstance(item, dict):
                continue
            reason = (
                (item.get("match_reason") or "").strip()
                or (item.get("compliance_reason") or "").strip()
                or "—"
            )
            _add_kv_card(
                doc,
                f"风险点 {i:02d}",
                [
                    ("原文语句", (item.get("text") or "—").strip() or "—"),
                    ("风险类型", _risk_tags(item)),
                    ("命中案例", _hit_case_line(item)),
                    ("处罚文号", _hit_doc_no(item)),
                    ("违法事实", _case_fact(item)),
                    ("匹配理由", reason),
                ],
            )

    # 三、人工复核建议
    _add_section_heading(doc, "三、人工复核建议")
    if human_review_done and (human_note or "").strip():
        _add_review_box(doc, human_note.strip(), placeholder=False)
    else:
        _add_review_box(
            doc,
            "供工作人员填写复核意见；无意见请填写「无」。"
            "（导出前请先在系统中完成人工复核，此处才会写入正式意见。）",
            placeholder=True,
        )

    # 页脚提示
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        "本报告由智能合规审查系统生成，结论供人工复核参考，不构成最终监管意见。"
    )
    _set_run_font(fr, size=8, color=MUTED)
    _para_spacing(foot, before=16, after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def default_export_filename(report: dict[str, Any]) -> str:
    mid = (report.get("material_id") or "")[:8] or "draft"
    src = report.get("file_name") or report.get("source_file") or "online"
    stem = _safe_filename(Path(str(src)).stem if src != "online" else "online_text")
    return f"合规审查报告_{stem}_{mid}.docx"
